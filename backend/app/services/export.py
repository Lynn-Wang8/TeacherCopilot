"""
DOCX 生成服务 — 将 ChapterData 导出为可编辑 Word 讲义

使用 python-docx 生成结构化教学文档：
  章节标题 → 题型分组 → 题目卡片（图片 + OCR + 易错点 + 模型 + 备注）
"""

import io
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from app.models.schemas import ChapterData, Question


def generate_docx(data: ChapterData) -> bytes:
    """
    根据章节数据生成 DOCX 文件

    Args:
        data: 完整的章节数据（按题型分组排好序）

    Returns:
        bytes: DOCX 文件二进制内容
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # ── 标题 ──
    title = doc.add_heading(f"{data.chapter.name} 错题整理", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # 空行

    # ── 按题型分组 ──
    groups: dict[str, list[Question]] = {}
    group_order: list[str] = []
    for q in data.questions:
        type_id = q.question_type.id
        if type_id not in groups:
            groups[type_id] = []
            group_order.append(type_id)
        groups[type_id].append(q)

    for type_id in group_order:
        qs = groups[type_id]
        type_name = qs[0].question_type.name

        # 题型标题
        doc.add_heading(type_name, level=1)

        for q in qs:
            # ── 题目卡片 ──
            # 题号
            doc.add_heading(f"题目 {q.question_id}", level=2)

            # 原题图片（如果有）
            if q.image and q.image.startswith("data:image"):
                try:
                    # 解析 base64 data URL
                    _, b64_data = q.image.split(",", 1)
                    img_bytes = base64.b64decode(b64_data)
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(4.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph("[图片加载失败]")

            # OCR 文字
            if q.ocr_text and not q.ocr_text.startswith("图片:"):
                p = doc.add_paragraph()
                run = p.add_run("OCR 文字：")
                run.bold = True
                run.font.size = Pt(10.5)
                p.add_run(q.ocr_text).font.size = Pt(10.5)

            # 易错点
            if q.common_mistake:
                p = doc.add_paragraph()
                run = p.add_run("易错点：")
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)  # 橙色
                p.add_run(q.common_mistake.name).font.size = Pt(10.5)
            else:
                p = doc.add_paragraph()
                run = p.add_run("易错点：")
                run.bold = True
                run.font.size = Pt(10.5)
                p.add_run("无").font.size = Pt(10.5)

            # 几何模型
            if q.geometry_model:
                p = doc.add_paragraph()
                run = p.add_run("几何模型：")
                run.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)  # 紫色
                p.add_run(q.geometry_model.name).font.size = Pt(10.5)
            else:
                p = doc.add_paragraph()
                run = p.add_run("几何模型：")
                run.bold = True
                run.font.size = Pt(10.5)
                p.add_run("无").font.size = Pt(10.5)

            # 教师备注
            if q.teacher_note:
                p = doc.add_paragraph()
                run = p.add_run("教师备注：")
                run.bold = True
                run.font.size = Pt(10.5)
                run.italic = True
                p.add_run(q.teacher_note).font.size = Pt(10.5)
            else:
                p = doc.add_paragraph()
                run = p.add_run("教师备注：")
                run.bold = True
                run.font.size = Pt(10.5)
                p.add_run("（未填写）").font.size = Pt(10.5)

            # 分隔线
            doc.add_paragraph("─" * 60)

    # ── 输出 ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
